import tkinter as tk
from tkinter import messagebox
from docx import Document
from datetime import datetime
import time
import sqlite3

# Создание или подключение к базе данных SQLite
conn = sqlite3.connect('payment_requests.db')
c = conn.cursor()

# Создание таблицы, если она не существует
c.execute('''CREATE TABLE IF NOT EXISTS payment_requests
             (id INTEGER PRIMARY KEY AUTOINCREMENT,
              employee_name TEXT,
              hours_worked REAL,
              hourly_rate REAL,
              supervisor_position TEXT,
              supervisor_initials TEXT,
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
conn.commit()

def create_payment_request(employee_name, hours_worked, hourly_rate, supervisor_position, supervisor_initials):
    try:
        # Валидация данных
        if hours_worked <= 0 or hourly_rate <= 0 or not employee_name or not supervisor_position or not supervisor_initials:
            messagebox.showerror("Ошибка", "Пожалуйста, заполните все поля корректно.")
            return

        # Создаем новый документ
        doc = Document()

        # Заголовок
        doc.add_heading('Заявление на почасовую оплату', 0)

        # Дата
        current_time = time.strftime("%d.%m.%Y")
        doc.add_paragraph(f'Дата: {current_time}')

        # Информация о сотруднике
        doc.add_paragraph(f'Сотрудник: {employee_name}')

        # Информация о почасовой оплате
        doc.add_paragraph(f'Количество отработанных часов: {hours_worked}')
        doc.add_paragraph(f'Почасовая ставка: {hourly_rate:.2f} руб.')

        # Общая сумма к оплате
        total_payment = hours_worked * hourly_rate
        doc.add_paragraph(f'Общая сумма к оплате: {total_payment:.2f} руб.')

        # Подпись сотрудника
        doc.add_paragraph(f'\nПодпись сотрудника: _____________________')

        # Подпись руководителя
        doc.add_paragraph(f'\n{supervisor_position}')
        doc.add_paragraph(f'{supervisor_initials}')
        doc.add_paragraph(f'Подпись: _____________________')

        # Сохранение документа
        filename = f'Заявление_{employee_name}_{datetime.now().strftime("%Y%m%d")}.docx'
        doc.save(filename)
        print(f'Заявление сохранено как {filename}')
        messagebox.showinfo('Успех', f'Заявление сохранено как {filename}')

        # Сохранение данных в базу
        c.execute('''INSERT INTO payment_requests (employee_name, hours_worked, hourly_rate, supervisor_position, supervisor_initials)
                     VALUES (?, ?, ?, ?, ?)''', (employee_name, hours_worked, hourly_rate, supervisor_position, supervisor_initials))
        conn.commit()

    except Exception as e:
        print(f'Ошибка при создании заявления: {e}')
        messagebox.showerror("Ошибка", "Произошла ошибка при создании заявления.")

def generate_report():
    try:
        # Запрос на получение данных о заявлениях из базы
        c.execute('''SELECT * FROM payment_requests''')
        rows = c.fetchall()

        if not rows:
            messagebox.showinfo("Отчет", "Нет данных о заявлениях.")
            return

        # Создание отчета
        report_doc = Document()
        report_doc.add_heading('Отчет о созданных заявлениях', 0)

        for row in rows:
            created_date = datetime.strptime(row[6], "%Y-%m-%d %H:%M:%S").strftime("%d.%m.%Y")
            report_doc.add_paragraph(f'Заявление от {created_date}')
            report_doc.add_paragraph(f'Сотрудник: {row[1]}')
            report_doc.add_paragraph(f'Отработанные часы: {row[2]}')
            report_doc.add_paragraph(f'Почасовая ставка: {row[3]}')
            report_doc.add_paragraph(f' {row[4]}')
            report_doc.add_paragraph(f' {row[5]}')
            report_doc.add_paragraph('-----------------------------------------')

        report_filename = f'Отчет_заявления_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        report_doc.save(report_filename)
        messagebox.showinfo('Отчет', f'Отчет сохранен как {report_filename}')

    except Exception as e:
        print(f'Ошибка при создании отчета: {e}')
        messagebox.showerror("Ошибка", "Произошла ошибка при создании отчета.")

def on_submit():
    try:
        employee_name = entry_employee_name.get()
        hours_worked = float(entry_hours_worked.get())
        hourly_rate = float(entry_hourly_rate.get())
        supervisor_position = entry_supervisor_position.get()
        supervisor_initials = entry_supervisor_initials.get()

        create_payment_request(employee_name, hours_worked, hourly_rate, supervisor_position, supervisor_initials)
    except ValueError:
        messagebox.showerror("Ошибка", "Пожалуйста, введите корректные значения для отработанных часов и почасовой ставки.")

# Создаем окно
root = tk.Tk()
root.title("Заявление на почасовую оплату")

# Поля ввода
tk.Label(root, text="Имя сотрудника").grid(row=0)
entry_employee_name = tk.Entry(root)
entry_employee_name.grid(row=0, column=1)

tk.Label(root, text="Отработанные часы").grid(row=1)
entry_hours_worked = tk.Entry(root)
entry_hours_worked.grid(row=1, column=1)

tk.Label(root, text="Почасовая ставка").grid(row=2)
entry_hourly_rate = tk.Entry(root)
entry_hourly_rate.grid(row=2, column=1)

tk.Label(root, text="Должность руководителя").grid(row=3)
entry_supervisor_position = tk.Entry(root)
entry_supervisor_position.grid(row=3, column=1)

tk.Label(root, text="Ф.И.О руководителя").grid(row=4)
entry_supervisor_initials = tk.Entry(root)
entry_supervisor_initials.grid(row=4, column=1)

# Кнопка отправки
submit_button = tk.Button(root, text="Создать заявление", command=on_submit)
submit_button.grid(row=5, columnspan=2)

# Кнопка для генерации отчета
report_button = tk.Button(root, text="Сгенерировать отчет", command=generate_report)
report_button.grid(row=6, columnspan=2)

# Запуск основного цикла
root.mainloop()

# Закрываем соединение с базой данных при выходе из программы
conn.close()
